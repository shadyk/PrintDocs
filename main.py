import os
import pandas as pd
from docx import Document
import tkinter as tk
from tkinter import ttk, messagebox, font
import platform
import subprocess
from datetime import datetime
from babel.dates import format_date

def convert_to_eastern_arabic(number):
    eastern_arabic_digits = {
        "0": "٠", "1": "١", "2": "٢", "3": "٣", "4": "٤",
        "5": "٥", "6": "٦", "7": "٧", "8": "٨", "9": "٩"
    }
    return "".join(eastern_arabic_digits[digit] for digit in str(number))

def replace_text_in_paragraph(paragraph, old_text, new_text):
    """Replace text even if it is split across multiple runs."""
    full_text = "".join(run.text for run in paragraph.runs)  # Combine all runs
    if old_text in full_text:
        full_text = full_text.replace(old_text, new_text)  # Replace text
        for run in paragraph.runs:
            run.text = ""  # Clear all runs
        paragraph.add_run(full_text)  # Add updated text as a single run

def open_directory(path):
    """Open the directory containing the generated file."""
    if platform.system() == "Windows":
        os.startfile(path)  # Open directory on Windows
    elif platform.system() == "Darwin":
        subprocess.Popen(["open", path])  # Open directory on macOS
    else:
        subprocess.Popen(["xdg-open", path])  # Open directory on Linux

# Function to fill the Word template
def fill_template(template_path, data_row, output_path):

    arabic_months = {
        "January": "كانون الثاني",
        "February": "شباط",
        "March": "آذار",
        "April": "نيسان",
        "May": "أيار",
        "June": "حزيران",
        "July": "تموز",
        "August": "آب",
        "September": "أيلول",
        "October": "تشرين الأول",
        "November": "تشرين الثاني",
        "December": "كانون الأول"
    }
            
    if not os.path.exists(template_path):
        messagebox.showerror("Error", f"Template file not found: {template_path}")
        return

    try:
        doc = Document(template_path)
    except Exception as e:
        messagebox.showerror("Error", f"Error loading document: {e}")
        return

    for key, value in data_row.items():
        key = key.strip()
        for paragraph in doc.paragraphs:
            if f'{{{key}}}' in paragraph.text:
                paragraph.text = paragraph.text.replace(f'{{{key}}}', str(value))
    
    today = datetime.today()

    day = convert_to_eastern_arabic(today.day)  # Convert day to Eastern Arabic
    month_name = arabic_months[today.strftime("%B")]
    year = convert_to_eastern_arabic(today.year)
    arabic_date = f"{day} {month_name} {year}" 

    for paragraph in doc.paragraphs:
        replace_text_in_paragraph(paragraph, "Today", arabic_date)

    
    doc.save(output_path)
    messagebox.showinfo("Success", f"Document saved to {output_path}")
    open_directory(os.path.dirname(output_path))  # Open the directory after saving

# Function to generate a document for the selected row
def generate_release_document():
    selected_index = listbox.curselection()
    if not selected_index:
        messagebox.showwarning("Warning", "Please select a row!")
        return
    selectedTemplate = release_situation_f
    selected_row = df.iloc[selected_index[0]]
    output_path = os.path.join(output_directory, f"اطلاق حال{selected_row.name + 1}.docx")
    for key, value in selected_row.items():
       key = key.strip()
       if key == 'Gender':
           if value == 'M':
               selectedTemplate = release_situation_m

    fill_template(selectedTemplate, selected_row, output_path)

def generate_baptisim_document():
    selected_index = listbox.curselection()
    if not selected_index:
        messagebox.showwarning("Warning", "Please select a row!")
        return
    selectedTemplate = baptisim_template_f
    selected_row = df.iloc[selected_index[0]]
    output_path = os.path.join(output_directory, f"معمودية{selected_row.name + 1}.docx")
    for key, value in selected_row.items():
       key = key.strip()
       if key == 'Gender':
           if value == 'M':
               selectedTemplate = baptisim_template_m
    fill_template(selectedTemplate, selected_row, output_path)

# Function to filter rows based on search input
def search_rows():
    search_term = search_var.get().strip().lower()
    if not search_term:
        listbox.delete(0, tk.END)
        for index, row in df.iterrows():
            listbox.insert(tk.END, " - ".join(map(str, row.values)))
        return

    filtered_rows = df.apply(lambda row: row.astype(str).str.contains(search_term, case=False).any(), axis=1)
    listbox.delete(0, tk.END)
    for index, row in df[filtered_rows].iterrows():
            listbox.insert(tk.END, " - ".join(map(str, row.values)))

path = ''
# path = 'Desktop/print/'

# Load Excel data
excel_file = path + 'data.xlsx'
baptisim_template_m = path + 'baptisim_template_m.docx'  
baptisim_template_f = path +'baptisim_template_f.docx'  
release_situation_m = path +'release_situation_m.docx'  
release_situation_f = path + 'release_situation_f.docx'  
output_directory = path + 'output_docs'

if not os.path.exists(excel_file):
    raise FileNotFoundError(f"Excel file not found: {excel_file}")

if not os.path.exists(output_directory):
    os.makedirs(output_directory)

df = pd.read_excel(excel_file, header=1)
df = df.loc[:, ~df.columns.str.contains('^Unnamed')]

# Create the Tkinter UI
root = tk.Tk()
root.title("اطلاق حال وورقات معموديات")

# Search bar
search_var = tk.StringVar()
search_entry = ttk.Entry(root, textvariable=search_var, width=50)
search_entry.pack(pady=10)
search_entry.bind("<KeyRelease>", lambda event: search_rows())

# Configure font for the Listbox
custom_font = font.Font(family="Traditional Arabic", size=20)  # Adjust size and font family as needed

# Create a frame to hold the Listbox and Scrollbar
frame = tk.Frame(root)
frame.pack(pady=10, fill=tk.BOTH, expand=True)

# Add a vertical scrollbar
scrollbar = ttk.Scrollbar(frame, orient=tk.VERTICAL)
scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

# Listbox to display rows
listbox = tk.Listbox(frame, width=100, height=20, font=custom_font, justify="right", yscrollcommand=scrollbar.set)
listbox.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

# Link the scrollbar to the Listbox
scrollbar.config(command=listbox.yview)

# Populate the listbox with all rows initially
for index, row in df.iterrows():
    listbox.insert(tk.END, " - ".join(map(str, row.values)))

# Create a frame for the buttons
button_frame = tk.Frame(root)
button_frame.pack(pady=10)

# Generate document buttons
generate_release_button = ttk.Button(button_frame, text="اطلاق حال", command=generate_release_document)
generate_release_button.pack(side=tk.LEFT, padx=5)

generate_baptisim_button = ttk.Button(button_frame, text="ورقة معمودية", command=generate_baptisim_document)
generate_baptisim_button.pack(side=tk.LEFT, padx=5)

# Run the Tkinter event loop
root.mainloop()